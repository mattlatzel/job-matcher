"""
CV Job Matcher — Backend
Run: python main.py
Requires: ANTHROPIC_API_KEY and JSEARCH_API_KEY set as environment variables
"""

import asyncio
import io
import json
import os
import re
import uuid
from pathlib import Path

import anthropic
import httpx
from fastapi import BackgroundTasks, FastAPI, File, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# ── App setup ────────────────────────────────────────────────────────────────

app = FastAPI(title="CV Job Matcher")

app.mount("/static", StaticFiles(directory="static"), name="static")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory session store  { session_id: { status, profile, results, total_jobs, error } }
sessions: dict = {}

ANTHROPIC_API_KEY  = os.getenv("ANTHROPIC_API_KEY")
JSEARCH_API_KEY    = os.getenv("JSEARCH_API_KEY")
ADZUNA_APP_ID      = os.getenv("ADZUNA_APP_ID")
ADZUNA_APP_KEY     = os.getenv("ADZUNA_APP_KEY")
APIFY_API_KEY      = os.getenv("APIFY_API_KEY")

SONNET = "claude-sonnet-5"   # profile extraction — needs quality reasoning
HAIKU  = "claude-haiku-4-5-20251001"   # job scoring — runs ~20+ times, needs speed


# ── CV Parsing ───────────────────────────────────────────────────────────────

def parse_pdf(content: bytes) -> str:
    import pdfplumber
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages)


def parse_docx(content: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(content))
    return "\n".join(p.text for p in doc.paragraphs)


# ── Claude: Quick title extraction (Haiku, ~0.5s) ────────────────────────────

async def quick_extract_title(client: anthropic.AsyncAnthropic, cv_text: str) -> str:
    resp = await client.messages.create(
        model=HAIKU,
        max_tokens=30,
        messages=[{"role": "user", "content":
            f"Extract only the job title from this CV. Return the title only, nothing else.\n\n{cv_text[:1500]}"}]
    )
    return next(b.text for b in resp.content if hasattr(b, "text")).strip()[:80]


# ── Claude: Extract structured profile from CV ───────────────────────────────

async def extract_profile(cv_text: str) -> dict:
    client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
    resp = await client.messages.create(
        model=SONNET,
        max_tokens=1024,
        messages=[{
            "role": "user",
            "content": f"""Analyze this CV and extract a structured profile. Return JSON only, no prose.

CV TEXT:
{cv_text[:8000]}

Return exactly this JSON structure:
{{
  "current_title": "most recent or target job title",
  "seniority": "junior | mid | senior | lead | executive",
  "years_experience": <number>,
  "core_domain": "1-sentence description of their main professional domain",
  "core_skills": ["top 6 most important skills"],
  "job_search_query": "best search query to find matching jobs (role + domain, e.g. 'Senior Product Manager SaaS')",
  "adjacent_titles": ["5-6 related job titles spanning BOTH lateral moves AND adjacent/broader roles. Cast wide — include different role types (e.g. BA, Solutions Consultant, PM) and related sectors. Better to include more than too few."],
  "target_sectors": ["4-5 sectors/industries where this person's skills are genuinely applicable, e.g. 'fintech', 'investment banking', 'capital markets', 'payments', 'insurtech'"],
  "adjacent_seniority": ["the 1-2 seniority levels directly above and below their current level, e.g. ['lead', 'principal'] for a senior person"],
  "location": "city/country or Remote",
  "summary": "2-sentence professional summary"
}}"""
        }]
    )
    text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    return json.loads(match.group() if match else text)


# ── JSearch: Fetch jobs ───────────────────────────────────────────────────────

async def _fetch_query(client: httpx.AsyncClient, query: str, pages: int = 3) -> list[dict]:
    """Fetch up to `pages` pages for a single query string."""
    jobs = []
    for page in range(1, pages + 1):
        try:
            resp = await client.get(
                "https://jsearch.p.rapidapi.com/search-v2",
                params={
                    "query": query,
                    "page": str(page),
                    "num_pages": "1",
                    "date_posted": "month",
                    "country": "gb",
                },
                headers={
                    "X-RapidAPI-Key": JSEARCH_API_KEY,
                    "X-RapidAPI-Host": "jsearch.p.rapidapi.com",
                },
            )
            data = resp.json()
            raw = data.get("data", [])
            batch = raw.get("jobs", []) if isinstance(raw, dict) else raw
            jobs.extend(batch)
            if not batch:
                break
        except Exception as e:
            print(f"  JSearch error ({query}, page {page}): {e}")
            break
    return jobs


async def _fetch_query_standalone(query: str, pages: int = 2) -> list[dict]:
    """Wrapper that creates its own httpx client — for parallel calls."""
    async with httpx.AsyncClient(timeout=30) as client:
        return await _fetch_query(client, query, pages)


async def _fetch_adzuna(keywords: str, location: str = "London", results_per_page: int = 50) -> list[dict]:
    """Fetch jobs from Adzuna API and normalise to our internal format."""
    if not ADZUNA_APP_ID or not ADZUNA_APP_KEY:
        return []
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            resp = await client.get(
                "https://api.adzuna.com/v1/api/jobs/gb/search/1",
                params={
                    "app_id":           ADZUNA_APP_ID,
                    "app_key":          ADZUNA_APP_KEY,
                    "what":             keywords,
                    "where":            location,
                    "distance":         15,
                    "results_per_page": results_per_page,
                    "content-type":     "application/json",
                },
            )
        data = resp.json()
        jobs = data.get("results", [])
        normalised = []
        for j in jobs:
            min_s = j.get("salary_min")
            max_s = j.get("salary_max")
            salary_str = None
            if min_s and max_s and min_s != max_s:
                salary_str = f"£{int(min_s):,} – £{int(max_s):,} / yr"
            elif max_s:
                salary_str = f"Up to £{int(max_s):,}"
            elif min_s:
                salary_str = f"From £{int(min_s):,}"

            normalised.append({
                "job_id":          f"adzuna_{j.get('id')}",
                "job_title":       j.get("title", ""),
                "employer_name":   j.get("company", {}).get("display_name", ""),
                "job_city":        j.get("location", {}).get("display_name", "London"),
                "job_country":     "United Kingdom",
                "job_is_remote":   False,
                "job_description": j.get("description", ""),
                "job_apply_link":  j.get("redirect_url", ""),
                "job_google_link": j.get("redirect_url", ""),
                "job_posted_at_datetime_utc": j.get("created", ""),
                "_salary_display": salary_str,
            })
        print(f"  Adzuna fetched {len(normalised)} jobs for '{keywords}'")
        return normalised
    except Exception as e:
        print(f"  Adzuna error ({keywords}): {e}")
        return []


async def _fetch_linkedin(search_queries: list[str], max_results: int = 40) -> list[dict]:
    """Fetch LinkedIn jobs via Apify curious_coder/linkedin-jobs-scraper."""
    if not APIFY_API_KEY:
        return []
    # Build LinkedIn search URLs — public job search (no login required)
    urls = []
    for q in search_queries[:4]:  # cap at 4 queries to control cost
        encoded = q.replace(" ", "%20")
        urls.append(f"https://www.linkedin.com/jobs/search/?keywords={encoded}&location=London%2C%20United%20Kingdom&f_TPR=r2592000")

    try:
        async with httpx.AsyncClient(timeout=120) as client:
            # Start the actor run
            run_resp = await client.post(
                "https://api.apify.com/v2/acts/curious_coder~linkedin-jobs-scraper/runs",
                headers={"Authorization": f"Bearer {APIFY_API_KEY}"},
                json={
                    "startUrls": urls,
                    "count": max_results,
                    "scrapeCompanyDetails": False,
                },
            )
            run_resp.raise_for_status()
            run_id = run_resp.json()["data"]["id"]
            print(f"  LinkedIn Apify run started: {run_id}")

            # Poll until finished (max 90s)
            for _ in range(18):
                await asyncio.sleep(5)
                status_resp = await client.get(
                    f"https://api.apify.com/v2/acts/curious_coder~linkedin-jobs-scraper/runs/{run_id}",
                    headers={"Authorization": f"Bearer {APIFY_API_KEY}"},
                )
                status = status_resp.json()["data"]["status"]
                if status in ("SUCCEEDED", "FAILED", "ABORTED", "TIMED-OUT"):
                    break

            if status != "SUCCEEDED":
                print(f"  LinkedIn Apify run ended with status: {status}")
                return []

            # Fetch results
            items_resp = await client.get(
                f"https://api.apify.com/v2/acts/curious_coder~linkedin-jobs-scraper/runs/{run_id}/dataset/items",
                headers={"Authorization": f"Bearer {APIFY_API_KEY}"},
                params={"limit": max_results},
            )
            items = items_resp.json()
            print(f"  LinkedIn Apify returned {len(items)} jobs")

            # Normalise to our internal job schema
            results = []
            for item in items:
                salary_text = item.get("salary") or ""
                results.append({
                    "job_id":          f"li_{item.get('id') or item.get('jobId', '')}",
                    "job_title":       item.get("title", ""),
                    "employer_name":   item.get("companyName", ""),
                    "job_city":        item.get("location", "London"),
                    "job_country":     "United Kingdom",
                    "job_is_remote":   "remote" in (item.get("location") or "").lower(),
                    "job_description": item.get("description") or item.get("descriptionText") or "",
                    "job_apply_link":  item.get("applyUrl") or item.get("url") or "#",
                    "job_posted_at_datetime_utc": item.get("postedAt") or "",
                    "_salary_display": salary_text or None,
                })
            return results

    except Exception as e:
        print(f"  LinkedIn Apify error: {e}")
        return []


async def fetch_jobs(profile: dict) -> list[dict]:
    title           = profile.get("current_title", "professional")
    domain          = profile.get("job_search_query") or title
    adjacent_titles   = profile.get("adjacent_titles", [])[:6]
    target_sectors    = profile.get("target_sectors", [])[:4]

    # JSearch queries — main + adjacent titles + sector-broadening
    jsearch_queries = [
        (f"{domain} in London",  6),
        (f"{title} in London",   4),
    ]
    for adj in adjacent_titles:
        jsearch_queries.append((f"{adj} in London", 3))
    for sector in target_sectors:
        jsearch_queries.append((f"{title} {sector} London", 2))

    print(f"  JSearch queries: {[q for q, _ in jsearch_queries]}")

    # Adzuna keywords — main + adjacent titles + sectors
    adzuna_keywords = [domain, title] + adjacent_titles[:4] + [f"{title} {s}" for s in target_sectors[:2]]

    # LinkedIn queries — main title + top adjacent titles
    linkedin_queries = [f"{title} London"] + [f"{adj} London" for adj in adjacent_titles[:3]]

    async with httpx.AsyncClient(timeout=45) as client:
        jsearch_results = await asyncio.gather(*[
            _fetch_query(client, q, pages=p) for q, p in jsearch_queries
        ])

    async def _linkedin_safe():
        if not APIFY_API_KEY:
            return []
        try:
            return await asyncio.wait_for(_fetch_linkedin(linkedin_queries), timeout=60)
        except Exception as e:
            print(f"  LinkedIn skipped: {e}")
            return []

    adzuna_results, linkedin_results = await asyncio.gather(
        asyncio.gather(*[_fetch_adzuna(kw) for kw in adzuna_keywords]),
        _linkedin_safe(),
    )

    # Merge and deduplicate by job_id AND by (normalised title + company)
    def _norm(s: str) -> str:
        s = (s or "").lower()
        s = re.sub(r'\b(senior|junior|lead|principal|staff|head of|director of|vp|associate)\b', '', s)
        return re.sub(r'[^a-z0-9]', '', s).strip()

    seen_ids: set    = set()
    seen_content: set = set()
    all_jobs: list[dict] = []
    for batch in list(jsearch_results) + list(adzuna_results) + [linkedin_results]:
        for job in batch:
            jid          = job.get("job_id")
            content_key  = (_norm(job.get("job_title", "")), _norm(job.get("employer_name", "")))
            if jid in seen_ids:
                continue
            if content_key[0] and content_key in seen_content:
                continue
            if jid:
                seen_ids.add(jid)
            if content_key[0]:
                seen_content.add(content_key)
            all_jobs.append(job)

    print(f"  Total unique jobs fetched: {len(all_jobs)}")

    # Build a human-readable search summary
    all_titles  = list(dict.fromkeys([title] + adjacent_titles))
    all_sectors = list(dict.fromkeys(target_sectors))

    titles_str  = ", ".join(f"**{t}**" for t in all_titles[:6])
    sectors_str = ", ".join(f"**{s}**" for s in all_sectors[:4])

    search_summary = f"I searched London job boards (Indeed, Adzuna and LinkedIn) for roles matching {titles_str}."
    if sectors_str:
        search_summary += f" I focused on {sectors_str} as your target sectors."
    search_summary += f" After removing duplicates I found **{len(all_jobs)} unique listings** to score against your CV."

    return all_jobs, search_summary


# ── Claude: Pre-filter jobs by title only ────────────────────────────────────

async def prefilter_jobs(
    client: anthropic.AsyncAnthropic,
    profile: dict,
    jobs: list[dict],
) -> list[dict]:
    """
    Quickly eliminate obvious mismatches using only job titles + companies.
    Returns only the jobs worth scoring in detail.
    """
    batch_size = 20
    passed: list[dict] = []

    seniority     = profile.get("seniority", "")
    adj_seniority = profile.get("adjacent_seniority", [])
    all_levels    = ", ".join(filter(None, [seniority] + adj_seniority)) or "mid"
    profile_line  = (
        f"{profile.get('current_title')} ({all_levels} level) "
        f"with {profile.get('years_experience')} years in {profile.get('core_domain')}"
    )

    for i in range(0, len(jobs), batch_size):
        batch = jobs[i : i + batch_size]
        titles_block = "\n".join(
            f"{j+1}. {job.get('job_title', 'N/A')} — {job.get('employer_name', 'N/A')}"
            for j, job in enumerate(batch)
        )

        resp = await client.messages.create(
            model=HAIKU,
            max_tokens=512,
            messages=[{
                "role": "user",
                "content": f"""You are a fast job screener. Decide which jobs are worth evaluating for this candidate.

CANDIDATE: {profile_line}

JOBS:
{titles_block}

Return a JSON array of "yes" or "no" for each job (in order).
"yes" = could plausibly be relevant — when in doubt, say yes.
"no"  = only if clearly wrong field or wildly wrong level (e.g. junior dev for a senior PM).

Be generous. It is better to include a borderline job than to miss a good match.

JSON array only. Example: ["yes","no","yes"]"""
            }]
        )

        text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
        match = re.search(r"\[.*?\]", text, re.DOTALL)
        decisions: list[str] = json.loads(match.group() if match else text)

        for j, job in enumerate(batch):
            if j < len(decisions) and decisions[j].lower() == "yes":
                passed.append(job)

    return passed


# ── Claude: Score a batch of jobs against the profile ────────────────────────

async def score_batch(
    client: anthropic.AsyncAnthropic,
    profile: dict,
    jobs: list[dict],
) -> list[dict]:
    jobs_block = ""
    for i, job in enumerate(jobs, 1):
        desc = (job.get("job_description") or "")[:3000]
        jobs_block += f"""
JOB {i}
Title: {job.get('job_title', 'N/A')}
Company: {job.get('employer_name', 'N/A')}
Location: {job.get('job_city', '')} {job.get('job_country', '')} {'(Remote)' if job.get('job_is_remote') else ''}
Description:
{desc}
────────────────────
"""

    profile_block = f"""Current title: {profile.get('current_title')}
Seniority: {profile.get('seniority')}
Years of experience: {profile.get('years_experience')}
Core domain: {profile.get('core_domain')}
Core skills: {', '.join(profile.get('core_skills', []))}
Summary: {profile.get('summary')}"""

    resp = await client.messages.create(
        model=SONNET,
        max_tokens=4096,
        messages=[{
            "role": "user",
            "content": f"""You are a senior recruiter evaluating job fit. Score each job and provide a structured breakdown.

Scoring rules:
- Reason about core experience vs core requirements — NOT keyword overlap.
- 90+: natural fit, strong alignment on seniority, domain, and skills.
- 70-89: good alignment but some gaps.
- Below 70: poor fit.
- Be honest. Do not inflate scores.

CANDIDATE:
{profile_block}

JOBS:
{jobs_block}

Return a JSON array (one object per job, in order):
[
  {{
    "score": <0-100>,
    "seniority_match": true | false,
    "domain_match": true | false,
    "strengths": ["1-3 specific strengths from the CV that match this job"],
    "gaps": ["1-3 specific gaps or missing requirements, or empty array if none"],
    "reason": "<2 sentences: overall fit summary>",
    "bullets": ["<what you'll own/build>", "<who they're looking for>", "<what makes this role distinct>"]
  }},
  ...
]

bullets must describe the ROLE itself (not the company background). Max 12 words each. JSON only. No prose."""
        }]
    )

    text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
    try:
        match = re.search(r"\[.*\]", text, re.DOTALL)
        scores: list[dict] = json.loads(match.group() if match else text)
    except Exception as e:
        print(f"  Score batch parse error: {e}")
        scores = []

    results = []
    for i, job in enumerate(jobs):
        s = scores[i] if i < len(scores) else {
            "score": 0, "seniority_match": False, "domain_match": False,
            "strengths": [], "gaps": [], "reason": "Could not evaluate."
        }

        # Build salary string — use pre-normalised Reed value if present
        salary = job.get("_salary_display")
        if not salary:
            min_s = job.get("job_min_salary")
            max_s = job.get("job_max_salary")
            currency = job.get("job_salary_currency") or "£"
            period = (job.get("job_salary_period") or "").lower()
            if min_s and max_s:
                salary = f"{currency}{int(min_s):,} – {currency}{int(max_s):,}"
                if period in ("year", "annual"):
                    salary += " / yr"
                elif period == "month":
                    salary += " / mo"
            elif min_s:
                salary = f"From {currency}{int(min_s):,}"
            elif max_s:
                salary = f"Up to {currency}{int(max_s):,}"

        results.append({
            "title":           job.get("job_title", "Unknown Role"),
            "company":         job.get("employer_name", "Unknown Company"),
            "location":        f"{job.get('job_city', '')} {job.get('job_country', '')}".strip() or "London",
            "is_remote":       bool(job.get("job_is_remote")),
            "apply_link":      job.get("job_apply_link") or job.get("job_google_link") or "#",
            "score":           int(s.get("score", 0)),
            "seniority_match": bool(s.get("seniority_match", False)),
            "domain_match":    bool(s.get("domain_match", False)),
            "strengths":       s.get("strengths", []),
            "gaps":            s.get("gaps", []),
            "reason":          s.get("reason", ""),
            "bullets":         s.get("bullets", []),
            "posted":          (job.get("job_posted_at_datetime_utc") or "")[:10],
            "salary":          salary,
            "description":     (job.get("job_description") or "")[:600].strip(),
        })
    return results


# ── Claude: Haiku quick-score (scores only, no breakdown) ────────────────────

async def quick_score_all(
    client: anthropic.AsyncAnthropic,
    profile: dict,
    jobs: list[dict],
) -> list[tuple]:
    """
    Fast Haiku pass returning (job, score) pairs sorted by score desc.
    Uses only title + company + 200-char description snippet.
    """
    profile_line = (
        f"{profile.get('current_title')}, {profile.get('years_experience')} yrs exp, "
        f"{profile.get('core_domain')}"
    )
    batch_size = 10
    scored: list[tuple] = []
    sem = asyncio.Semaphore(8)

    async def score_quick_batch(batch):
        jobs_text = "\n".join(
            f"{j+1}. {job.get('job_title','?')} @ {job.get('employer_name','?')}: "
            f"{(job.get('job_description') or '')[:200]}"
            for j, job in enumerate(batch)
        )
        async with sem:
            try:
                resp = await client.messages.create(
                    model=HAIKU,
                    max_tokens=80,
                    messages=[{"role": "user", "content":
                        f"Candidate: {profile_line}\n\nRate each job 0-100 for fit. "
                        f"JSON array of integers only.\n\n{jobs_text}\n\nScores:"}]
                )
                text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
                match = re.search(r"\[.*?\]", text, re.DOTALL)
                scores = json.loads(match.group() if match else "[]")
            except Exception as e:
                print(f"  Quick score error: {e}")
                scores = []
            for j, job in enumerate(batch):
                scored.append((job, int(scores[j]) if j < len(scores) else 50))

    batches = [jobs[i:i+batch_size] for i in range(0, len(jobs), batch_size)]
    await asyncio.gather(*[score_quick_batch(b) for b in batches])
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored


# ── Claude: Sonnet deep-score with streaming to session ──────────────────────

async def score_and_stream(
    session: dict,
    client: anthropic.AsyncAnthropic,
    profile: dict,
    jobs: list[dict],
) -> None:
    """Score jobs with Sonnet, appending results to session as each batch finishes."""
    session["scoring_total"] = len(jobs)
    session["scoring_done"]  = 0
    batches = [jobs[i:i+3] for i in range(0, len(jobs), 3)]
    sem = asyncio.Semaphore(4)

    async def do_batch(batch):
        async with sem:
            results = await score_batch(client, profile, batch)
            session["results"].extend(results)
            session["results"].sort(key=lambda x: x["score"], reverse=True)
            session["scoring_done"] += len(results)

    await asyncio.gather(*[do_batch(b) for b in batches])


# ── Claude: CV gap analysis across all scored jobs ────────────────────────────

async def analyze_cv_gaps(
    client: anthropic.AsyncAnthropic,
    profile: dict,
    scored_jobs: list[dict],
) -> list[dict]:
    """
    Meta-analysis across top scored jobs to identify recurring CV gaps.
    Returns a list of {gap, suggestion, frequency} objects.
    """
    top = scored_jobs[:20]  # look at top 20 matched jobs

    jobs_gaps = ""
    for job in top:
        gaps = job.get("gaps", [])
        if gaps:
            jobs_gaps += f"- {job['title']} ({job['score']}% match): {'; '.join(gaps)}\n"

    if not jobs_gaps:
        return []

    profile_block = f"""Title: {profile.get('current_title')}
Seniority: {profile.get('seniority')}
Years of experience: {profile.get('years_experience')}
Core domain: {profile.get('core_domain')}
Core skills: {', '.join(profile.get('core_skills', []))}
Summary: {profile.get('summary')}"""

    resp = await client.messages.create(
        model=SONNET,
        max_tokens=2048,
        messages=[{
            "role": "user",
            "content": f"""You are a career coach analysing a candidate's CV against London job market requirements.

CANDIDATE PROFILE:
{profile_block}

GAPS IDENTIFIED ACROSS MATCHED JOBS:
{jobs_gaps}

Based on these recurring gaps, identify the 3-5 most impactful improvements this candidate should make to their CV to increase their match rate.

Be specific and actionable. Focus on gaps that appear multiple times — these are the most valuable to address.

Return the 3 most important gaps only. Be very brief — a few words each.

Return JSON array only:
[
  {{
    "gap": "3-5 word gap title",
    "action": "what to add or fix in CV (max 10 words)"
  }},
  ...
]"""
        }]
    )

    text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
    try:
        match = re.search(r"\[.*\]", text, re.DOTALL)
        return json.loads(match.group() if match else text)
    except Exception as e:
        print(f"  Gap analysis parse error: {e}\n  Raw: {text[:200]}")
        return []


# ── Chat: conversational CV intake ───────────────────────────────────────────

CHAT_SYSTEM = """You are Chelsea — a sharp, warm, and relentlessly driven career advisor who lives to find people the job they deserve.

Your personality: warm but direct, a little cheeky, always encouraging. You make the candidate feel like you're fighting their corner.

Writing style:
- Use **bold** for company names, role titles, and key skills (e.g. **Lucera**, **Senior PM**, **FIX connectivity**)
- Keep messages SHORT. 2-4 sentences maximum per message.
- Never write walls of text. Break thoughts into short, punchy paragraphs.
- Never use dashes (em dash, en dash, or hyphens) as punctuation. Use commas, periods, or semicolons instead.

Have a natural conversation to uncover:
1. What they actually built or owned in their key roles (CVs are always vague; push past the jargon)
2. Whether they lean more technical/hands-on or strategic/managerial
3. What they specifically want next: sector, role type, company stage, remote vs office
4. Any strong preferences or hard constraints (salary, culture, things they would never do again)

Rules:
- Ask ONE focused question at a time; punchy and conversational
- Always reference something specific from their CV; make it feel personal, not like a form
- Do NOT ask generic questions like "tell me about yourself"
- After the candidate has answered 4 or more questions, wrap up naturally
- In your wrap-up, use bold to summarise what you've learned in 2-3 short bullets, then say you are starting the search
- At the very end of your wrap-up message ONLY, append the exact text: |||DONE|||

CV:
{cv_text}"""


async def chat_turn(
    client: anthropic.AsyncAnthropic,
    cv_text: str,
    messages: list[dict],
) -> dict:
    user_count = sum(1 for m in messages if m["role"] == "user")

    system = CHAT_SYSTEM.format(cv_text=cv_text[:5000])

    # If no messages yet, bootstrap with a user opener
    msgs = messages if messages else [{"role": "user", "content": "Hi, I just uploaded my CV."}]

    # After 4 user answers, explicitly instruct Claude to wrap up
    if user_count >= 4:
        system += "\n\nIMPORTANT: You now have enough information. This must be your final message. Wrap up warmly, tell them you're starting the job search, and end with |||DONE|||"

    resp = await client.messages.create(
        model=SONNET,
        max_tokens=350,
        system=system,
        messages=msgs,
    )

    text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
    done = "|||DONE|||" in text
    clean = text.replace("|||DONE|||", "").strip()
    return {"message": clean, "done": done}


REFINE_SYSTEM = """You are Chelsea — a sharp, warm career advisor sitting next to the candidate, looking at their job results together.

You're ready to:
1. React to jobs they open — like a recruiter looking over their shoulder
2. Answer questions about specific roles, companies, or fit
3. Refine and re-run the search if needed

Your personality: warm, direct, a little cheeky, always in the candidate's corner.

Writing style:
- Use **bold** for company names, role titles, and key terms
- Keep replies SHORT. 2-3 sentences maximum.
- Never use dashes as punctuation; use commas, periods, or semicolons instead.
- Sound like a human recruiter, not a report.

Special trigger: when the user message starts with [JOB_OPENED], they just clicked open a job card. React like a recruiter sitting next to them who just spotted something interesting. Lead with your gut read on the fit in ONE sentence, then one specific reason why it stands out (or a honest flag if something's off). Keep it to 2 sentences total. Do not start with "Oh" or "Ooh" every time, vary your opener.

If they want a refined search, ask at most ONE clarifying question if truly needed, then confirm what you're doing and end with |||DONE||| on its own line. Do NOT output |||DONE||| for general chat.

CV:
{cv_text}"""


async def refine_turn(
    client: anthropic.AsyncAnthropic,
    cv_text: str,
    messages: list[dict],
    active_job: dict | None = None,
) -> dict:
    system = REFINE_SYSTEM.format(cv_text=cv_text[:4000])
    if active_job:
        strengths = ", ".join(active_job.get("strengths", []))
        gaps      = ", ".join(active_job.get("gaps", []))
        system += f"""

The candidate currently has this job expanded on their screen:
- Title: {active_job.get("title", "")}
- Company: {active_job.get("company", "")}
- Match score: {active_job.get("score", "")}%
- Salary: {active_job.get("salary_display") or active_job.get("_salary_display", "not listed")}
- Seniority match: {active_job.get("seniority_match", "")}
- Domain match: {active_job.get("domain_match", "")}
- Strengths: {strengths}
- Gaps: {gaps}
- Description: {(active_job.get("description") or "")[:300]}

If they ask about this role, use these details to explain the score and fit precisely."""
    resp = await client.messages.create(
        model=SONNET,
        max_tokens=400,
        system=system,
        messages=messages,
    )
    text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
    done = "|||DONE|||" in text
    clean = text.replace("|||DONE|||", "").strip()
    return {"message": clean, "done": done}


async def extract_profile_from_conversation(
    cv_text: str,
    messages: list[dict],
) -> dict:
    """Richer profile extraction using CV + what candidate told us in chat."""
    # Build readable conversation (skip bootstrap message)
    convo_lines = []
    for m in messages[1:]:  # skip bootstrap
        role = "Advisor" if m["role"] == "assistant" else "Candidate"
        convo_lines.append(f"{role}: {m['content']}")
    convo = "\n".join(convo_lines)

    client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
    resp = await client.messages.create(
        model=SONNET,
        max_tokens=1500,
        messages=[{
            "role": "user",
            "content": f"""Extract a rich professional profile from this CV and follow-up conversation.
The conversation reveals what the candidate actually did and what they're looking for.

CV:
{cv_text[:6000]}

CONVERSATION:
{convo}

Return exactly this JSON structure:
{{
  "current_title": "most recent or target job title",
  "seniority": "junior | mid | senior | lead | executive",
  "years_experience": <number>,
  "core_domain": "1-sentence description of their main professional domain",
  "core_skills": ["top 6 skills, informed by the conversation"],
  "job_search_query": "best search query to find matching jobs",
  "adjacent_titles": ["5-6 related job titles spanning BOTH lateral moves AND adjacent/broader roles. Cast wide — include different role types and related sectors. Prioritise titles the candidate expressed interest in during the conversation."],
  "target_sectors": ["4-5 sectors/industries where this person's skills are genuinely applicable. Informed by what the candidate said they want."],
  "adjacent_seniority": ["the 1-2 seniority levels directly above and below their current level"],
  "what_they_want": "1-2 sentences on what they specifically said they want next",
  "location": "city/country or Remote",
  "summary": "2-sentence professional summary that incorporates what they told us"
}}

JSON only. No prose."""
        }]
    )

    text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
    match = re.search(r"\{.*\}", text, re.DOTALL)
    return json.loads(match.group() if match else text)


# ── Claude: Salary benchmarking ──────────────────────────────────────────────

async def analyze_salary(
    client: anthropic.AsyncAnthropic,
    profile: dict,
    scored_jobs: list[dict],
) -> str | None:
    jobs_with_salary = [j for j in scored_jobs if j.get("salary") and j.get("score", 0) >= 60][:20]
    if len(jobs_with_salary) < 3:
        return None

    salary_lines = "\n".join(
        f"- {j['title']} at {j['company']}: {j['salary']}"
        for j in jobs_with_salary
    )

    resp = await client.messages.create(
        model=HAIKU,
        max_tokens=80,
        messages=[{"role": "user", "content": f"""You are Chelsea, a career advisor. Based on these live London job listings for a {profile.get('current_title')} with {profile.get('years_experience')} years experience:

{salary_lines}

Write ONE warm, concise sentence (max 25 words) giving the salary range you are seeing. Start with "Based on what I'm seeing..." and do not use dashes."""}]
    )
    return next(b.text for b in resp.content if hasattr(b, "text")).strip()


# ── Background worker ─────────────────────────────────────────────────────────

async def process_cv(session_id: str, cv_text: str, conversation_messages: list = None) -> None:
    session = sessions[session_id]
    try:
        ai_client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

        # 1. Extract profile (enriched if we have conversation)
        print("▶ Step 1: Extracting profile…")
        session["status"] = "extracting_profile"
        if conversation_messages:
            profile = await extract_profile_from_conversation(cv_text, conversation_messages)
        else:
            profile = await extract_profile(cv_text)
        session["profile"] = profile
        print(f"✓ Profile: {profile.get('current_title')}")

        # 2. Fetch jobs
        print("▶ Step 2: Fetching jobs…")
        session["status"] = "fetching_jobs"
        jobs, search_summary = await fetch_jobs(profile)
        session["total_jobs"] = len(jobs)
        session["search_summary"] = search_summary
        print(f"✓ Fetched {len(jobs)} jobs")

        # 3. Pre-filter by title
        print("▶ Step 3: Pre-filtering…")
        session["status"] = "scoring"
        jobs = await prefilter_jobs(ai_client, profile, jobs)
        print(f"✓ Pre-filter passed: {len(jobs)} jobs")

        # 4. Sonnet deep scoring
        print("▶ Step 4: Deep scoring…")
        batches = [jobs[i:i+3] for i in range(0, len(jobs), 3)]
        sem = asyncio.Semaphore(6)

        async def score_with_sem(batch):
            async with sem:
                return await score_batch(ai_client, profile, batch)

        batch_results = await asyncio.gather(*[score_with_sem(b) for b in batches])
        all_results = [r for br in batch_results for r in br]
        all_results.sort(key=lambda x: x["score"], reverse=True)
        session["results"] = all_results
        print(f"✓ Scored {len(all_results)} jobs")

        # 5. Gap analysis + salary insight (parallel)
        print("▶ Step 5: Gap analysis + salary benchmarking…")
        gap_task    = analyze_cv_gaps(ai_client, profile, all_results)
        salary_task = analyze_salary(ai_client, profile, all_results)
        session["gap_analysis"], session["salary_insight"] = await asyncio.gather(gap_task, salary_task)
        print(f"✓ Found {len(session['gap_analysis'])} gaps | salary insight: {bool(session['salary_insight'])}")

        session["status"] = "done"
        session["chat_status"] = "done"
        print(f"✓ Done! {len([r for r in all_results if r['score'] >= 70])} jobs ≥70%")

    except Exception as exc:
        print(f"✗ ERROR: {exc}")
        import traceback; traceback.print_exc()
        session["status"] = "error"
        session["error"] = str(exc)


# ── API Routes ────────────────────────────────────────────────────────────────

class ChatMessageBody(BaseModel):
    message: str
    active_job: dict | None = None

@app.post("/api/chat/start")
async def chat_start(file: UploadFile = File(...)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(500, "ANTHROPIC_API_KEY not set")
    if not JSEARCH_API_KEY:
        raise HTTPException(500, "JSEARCH_API_KEY not set")

    content = await file.read()
    name = (file.filename or "").lower()

    if name.endswith(".pdf"):
        cv_text = parse_pdf(content)
    elif name.endswith(".docx"):
        cv_text = parse_docx(content)
    elif name.endswith(".txt"):
        cv_text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(400, "Unsupported file type. Please upload PDF, DOCX, or TXT.")

    if len(cv_text.strip()) < 100:
        raise HTTPException(400, "Could not extract text from the CV. Try a different format.")

    session_id = str(uuid.uuid4())
    ai_client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
    first = await chat_turn(ai_client, cv_text, [])

    bootstrap = [{"role": "user", "content": "Hi, I just uploaded my CV."}]
    sessions[session_id] = {
        "chat_status": "chatting",
        "cv_text":     cv_text,
        "messages":    bootstrap + [{"role": "assistant", "content": first["message"]}],
        "docx_bytes":  content if name.endswith(".docx") else None,
        # search fields
        "status":      "starting",
        "profile":     None,
        "results":     [],
        "gap_analysis":   [],
        "salary_insight": None,
        "total_jobs":     0,
        "error":          None,
    }

    return {
        "session_id": session_id,
        "message":    first["message"],
        "done":       first["done"],
        "is_docx":    name.endswith(".docx"),
    }


@app.post("/api/chat/message/{session_id}")
async def chat_message_endpoint(
    session_id: str,
    body: ChatMessageBody,
    background_tasks: BackgroundTasks,
):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    s = sessions[session_id]
    if s.get("chat_status") not in ("chatting", "done"):
        raise HTTPException(400, "Chat is not active")

    s["messages"].append({"role": "user", "content": body.message})

    ai_client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)

    if s.get("chat_status") == "done":
        result = await refine_turn(ai_client, s["cv_text"], s["messages"], body.active_job)
    else:
        result = await chat_turn(ai_client, s["cv_text"], s["messages"])

    s["messages"].append({"role": "assistant", "content": result["message"]})

    if result["done"]:
        s["chat_status"] = "searching"
        background_tasks.add_task(process_cv, session_id, s["cv_text"], s["messages"])

    return {"message": result["message"], "done": result["done"]}


@app.post("/api/analyze")
async def analyze(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    if not ANTHROPIC_API_KEY:
        raise HTTPException(500, "ANTHROPIC_API_KEY not set")
    if not JSEARCH_API_KEY:
        raise HTTPException(500, "JSEARCH_API_KEY not set")

    content = await file.read()
    name = (file.filename or "").lower()

    if name.endswith(".pdf"):
        cv_text = parse_pdf(content)
    elif name.endswith(".docx"):
        cv_text = parse_docx(content)
    elif name.endswith(".txt"):
        cv_text = content.decode("utf-8", errors="ignore")
    else:
        raise HTTPException(400, "Unsupported file type. Please upload PDF, DOCX, or TXT.")

    if len(cv_text.strip()) < 100:
        raise HTTPException(400, "Could not extract text from the CV. Try a different format.")

    session_id = str(uuid.uuid4())
    sessions[session_id] = {
        "status":       "starting",
        "profile":      None,
        "results":      [],
        "gap_analysis": [],
        "total_jobs":   0,
        "error":        None,
        "docx_bytes":   content if name.endswith(".docx") else None,
    }
    background_tasks.add_task(process_cv, session_id, cv_text)
    return {"session_id": session_id, "is_docx": name.endswith(".docx")}


@app.get("/api/status/{session_id}")
async def get_status(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    s = sessions[session_id]
    return {
        "status":        s["status"],
        "total_jobs":    s["total_jobs"],
        "profile":       s["profile"],
        "error":         s["error"],
    }


@app.get("/api/results/{session_id}")
async def get_results(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    s = sessions[session_id]
    return {
        "jobs":           s.get("results", []),
        "gap_analysis":   s.get("gap_analysis", []),
        "salary_insight": s.get("salary_insight"),
        "search_summary": s.get("search_summary", ""),
    }


# ── Claude: Improve CV based on gap analysis ─────────────────────────────────

async def improve_cv_content(
    client: anthropic.AsyncAnthropic,
    docx_bytes: bytes,
    gap_analysis: list[dict],
    profile: dict,
) -> bytes:
    import docx as docx_lib

    doc = docx_lib.Document(io.BytesIO(docx_bytes))

    # Extract non-empty paragraphs with their index
    para_lines = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            para_lines.append(f"{i}: {text}")

    paragraphs_block = "\n".join(para_lines)

    gaps_block = "\n".join(
        f"- Gap: {g.get('gap')} | Action: {g.get('action')}" for g in gap_analysis
    )

    resp = await client.messages.create(
        model=SONNET,
        max_tokens=4096,
        messages=[{
            "role": "user",
            "content": f"""You are a professional CV editor. Your task is to improve specific lines in this CV to address identified skill gaps.

CANDIDATE PROFILE:
Title: {profile.get('current_title')}
Domain: {profile.get('core_domain')}
Skills: {', '.join(profile.get('core_skills', []))}

GAPS TO ADDRESS:
{gaps_block}

CV PARAGRAPHS (format: index: text):
{paragraphs_block}

Instructions:
- Identify which paragraphs should be rewritten to address the gaps above.
- Only change paragraphs where an improvement is genuinely needed — do NOT change headings, contact info, dates, or sections unrelated to the gaps.
- Rewrite the selected paragraphs to naturally incorporate the missing skills or experience framing.
- Keep the same general length and tone. Do not invent experience that isn't there — reframe and strengthen what exists.
- Return ONLY the paragraphs you changed.

Return JSON array only:
[
  {{"index": <paragraph_index>, "new_text": "improved paragraph text"}},
  ...
]"""
        }]
    )

    text = next(b.text for b in resp.content if hasattr(b, "text")).strip()
    try:
        match = re.search(r"\[.*\]", text, re.DOTALL)
        changes: list[dict] = json.loads(match.group() if match else text)
    except Exception as e:
        print(f"  CV improve parse error: {e}")
        changes = []

    print(f"  Applying {len(changes)} paragraph changes to CV")

    for change in changes:
        idx = change.get("index")
        new_text = change.get("new_text", "").strip()
        if idx is None or not new_text:
            continue
        if idx >= len(doc.paragraphs):
            continue
        para = doc.paragraphs[idx]
        if not para.runs:
            # No runs — just set text directly (rare)
            para.clear()
            para.add_run(new_text)
        else:
            # Preserve first run's character formatting, wipe the rest
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read()


@app.post("/api/improve-cv/{session_id}")
async def improve_cv(session_id: str):
    if session_id not in sessions:
        raise HTTPException(404, "Session not found")
    s = sessions[session_id]
    if s.get("status") != "done":
        raise HTTPException(400, "Analysis not complete yet")
    if not s.get("docx_bytes"):
        raise HTTPException(400, "No Word document in this session")

    client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
    improved_bytes = await improve_cv_content(
        client,
        s["docx_bytes"],
        s.get("gap_analysis", []),
        s.get("profile", {}),
    )

    return StreamingResponse(
        io.BytesIO(improved_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=improved_cv.docx"},
    )


@app.get("/")
async def serve_frontend():
    html = (Path(__file__).parent / "index.html").read_text()
    return HTMLResponse(content=html)


# ── Entrypoint ────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    print(f"\n🚀  CV Job Matcher running at http://localhost:{port}\n")
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
